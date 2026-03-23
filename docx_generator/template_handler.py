"""DOCX template renderer for plantilla DPI Canarias.

INVARIANT: templates/plantilla.docx is NEVER modified.
render_template() always reads the original and writes a NEW file to outputs/.
"""

import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from config import BASE_DIR, TEMPLATES_DIR
from scoring_engine import CRITERION_OPTIONS, SCORE_MAP

OUTPUT_DIR = BASE_DIR / "outputs"
DEFAULT_TEMPLATE = TEMPLATES_DIR / "plantilla.docx"

# Markers en párrafos libres → clave en free_texts
FREE_TEXT_MARKERS = {
    "[Poner texto]": "definicion_potencial",
    "[Redactar]": "conclusiones",
}

# Claves DAFO que aparecen como placeholders {{}} en la plantilla (fallback)
DAFO_KEYS = {
    "dafo_debilidades",
    "dafo_amenazas",
    "dafo_fortalezas",
    "dafo_oportunidades",
}

# BUG 3 fix: positional mapping (row_idx, col_idx) → free_text key
_DAFO_POSITIONS: dict[tuple[int, int], str] = {
    (2, 0): "dafo_debilidades",
    (2, 1): "dafo_amenazas",
    (3, 0): "dafo_fortalezas",
    (3, 1): "dafo_oportunidades",
}

_DAFO_HEADER_KEYWORDS = {
    "dafo",
    "debilidades",
    "fortalezas",
    "amenazas",
    "oportunidades",
}

# Positional row indices in doc.tables[1] per DPI criterion.
# col[0] = option label, col[1] = [Valorar] / {{placeholder}} to clear.
_OPTION_MAP: dict[str, list[int]] = {
    "situacion_empresa": [7, 8, 9],
    "num_empleados": [12, 13],
    "facturacion": [16, 17, 18, 19],
    "evolucion_facturacion": [22, 23, 24],
    "recursos_economicos": [27, 28],
    "experiencia_internacional": [33, 34, 35],
    "alcance_actividad": [38, 39, 40],
    "num_paises": [43, 44, 45],
    "personal_internacionalizacion": [48, 49],
    "involuccion_gerencia": [52, 53, 54, 55],
    "adaptacion_demanda": [58, 59, 60],
    "adaptacion_producto": [63, 64, 65],
    "tiene_web": [70, 71],
    "ecommerce": [74, 75, 76],
    "mercados_electronicos": [79, 80, 81, 82],
    "redes_sociales": [85, 86, 87],
}

# Indicator row = min(option_rows) - 1 for each criterion.
# That row's col[1] already has #FFC000 background in the template.
_INDICATOR_ROW_MAP: dict[str, int] = {
    crit: min(rows) - 1 for crit, rows in _OPTION_MAP.items()
}


def _write_score_values(doc: Document, selections: dict[str, str | None]) -> None:
    """Write the numeric score for each criterion into its pre-formatted golden cell.

    The template already has #FFC000 background on table[1].rows[indicator_row].cells[1].
    We only write the integer score string — never change background or font.
    """
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    rows = table.rows
    n_rows = len(rows)

    for criterion, indicator_row in _INDICATOR_ROW_MAP.items():
        if indicator_row >= n_rows:
            continue
        cells = rows[indicator_row].cells
        if len(cells) < 2:
            continue
        selected = selections.get(criterion)
        score = SCORE_MAP.get(criterion, {}).get(selected, 0) if selected else 0
        _set_cell_text(cells[1], str(score))


# Criteria that belong to each template block (by visual position in the DOCX,
# independent of the scoring_matrix.json block assignment).
# Row 5  = Bloque 1 subtotal, Row 31 = Bloque 2, Row 68 = Bloque 3.
_BLOCK_SUBTOTAL_ROWS: dict[int, list[str]] = {
    5: ["situacion_empresa", "num_empleados", "facturacion", "evolucion_facturacion"],
    31: [
        "experiencia_internacional",
        "alcance_actividad",
        "num_paises",
        "involuccion_gerencia",
        "adaptacion_demanda",
        "adaptacion_producto",
    ],
    68: ["tiene_web", "ecommerce", "mercados_electronicos", "redes_sociales"],
}


def _write_totals(doc: Document, selections: dict[str, str | None]) -> None:
    """Write block subtotals (rows 5, 31, 68) and DPI total (rows 0-1) into table[1]."""
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    rows = table.rows
    n_rows = len(rows)

    def _score(criterion: str) -> int:
        val = selections.get(criterion)
        return SCORE_MAP.get(criterion, {}).get(val, 0) if val else 0

    grand_total = 0
    for subtotal_row, criteria in _BLOCK_SUBTOTAL_ROWS.items():
        block_sum = sum(_score(c) for c in criteria)
        grand_total += block_sum
        if subtotal_row < n_rows and len(rows[subtotal_row].cells) >= 2:
            _set_cell_text(rows[subtotal_row].cells[1], str(block_sum))

    # DPI total into the two header rows (0 and 1)
    for row_idx in (0, 1):
        if row_idx < n_rows and len(rows[row_idx].cells) >= 2:
            _set_cell_text(rows[row_idx].cells[1], str(grand_total))


# Placeholder keys in table[1] header rows that must be cleared to "".
_DPI_TABLE_PLACEHOLDERS = (
    "VALOR_CONSTITUCION",
    "valor_empleados",
    "valor_facturacion",
    "evolucion_facturacion",
    "evolucion_recursos",
    "valorar",
)


# ─── Helpers ──────────────────────────────────────────────────────────────────


def _normalize_text(text: str) -> str:
    """Strip, lowercase, and remove combining accent marks for robust Spanish comparison."""
    nfkd = unicodedata.normalize("NFKD", text.strip().lower())
    return "".join(c for c in nfkd if not unicodedata.combining(c))


def _find_template(document_type: str = "default") -> Path:
    """Locate the right template file for the document type."""
    specific = TEMPLATES_DIR / f"{document_type}.docx"
    if specific.exists():
        return specific
    if DEFAULT_TEMPLATE.exists():
        return DEFAULT_TEMPLATE
    raise FileNotFoundError(
        f"No template found in {TEMPLATES_DIR}. "
        "Upload a 'plantilla.docx' to /root/autoreporte/templates/"
    )


def _set_cell_text(cell, text: str) -> None:
    """Write text into a cell's first paragraph, collapsing runs."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    p = paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(text)
    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ""


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Replace {{key}} placeholders with case-insensitive key lookup."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # BUG 1 fix: case-insensitive key matching
    lower_map = {k.lower(): v for k, v in replacements.items()}
    new_text = full_text
    for m in re.finditer(r"\{\{([^}]+)\}\}", full_text):
        value = lower_map.get(m.group(1).lower())
        if value is not None:
            new_text = new_text.replace(m.group(0), value)

    if new_text != full_text and paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_marker_in_paragraph(paragraph, free_texts: dict[str, str]) -> None:
    """Replace [Poner texto] / [Redactar] markers with free text content.

    Replaces the *entire* paragraph text with the replacement value to prevent
    heading-text concatenation when the marker is embedded in a labelled paragraph.
    """
    full_text = paragraph.text
    for marker, key in FREE_TEXT_MARKERS.items():
        if marker in full_text and key in free_texts:
            replacement = free_texts[key] or ""
            if paragraph.runs:
                paragraph.runs[0].text = replacement
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(replacement)
            return  # one marker per paragraph is enough


def _iter_all_paragraphs(doc: Document):
    """Yield every paragraph in the document: body + all table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _apply_direct_fields_to_paragraphs(
    doc: Document, direct_fields: dict[str, str]
) -> None:
    """Replace {{placeholder}} in all paragraphs (body + table cells)."""
    for paragraph in _iter_all_paragraphs(doc):
        _replace_in_paragraph(paragraph, direct_fields)


def _apply_free_texts_to_paragraphs(doc: Document, free_texts: dict[str, str]) -> None:
    """Replace [Poner texto]/[Redactar] markers and {{dafo_*}} fallback placeholders."""
    dafo_replacements = {k: free_texts.get(k, "") for k in DAFO_KEYS}
    for paragraph in _iter_all_paragraphs(doc):
        _replace_marker_in_paragraph(paragraph, free_texts)  # BUG 4 fix
        _replace_in_paragraph(paragraph, dafo_replacements)  # {{dafo_*}} fallback


def _clean_corrupted_headings(doc: Document) -> None:
    """Remove text that leaked into section headings from previous marker replacements.

    The 'Conclusiones de análisis' heading may have accumulated extra text
    (e.g. 'egia en tres pilares inmediatos:') if a previous run embedded
    [Redactar] inside the heading paragraph. This pass resets it to the
    canonical heading text.
    """
    _HEADING_CANONICAL: dict[str, str] = {
        "conclusiones de analisis": "Conclusiones de análisis",
    }
    for paragraph in doc.paragraphs:
        if "Heading" not in paragraph.style.name:
            continue
        norm = unicodedata.normalize("NFKD", paragraph.text.strip().lower())
        norm = "".join(c for c in norm if not unicodedata.combining(c))
        for prefix, canonical in _HEADING_CANONICAL.items():
            if norm.startswith(prefix) and norm != prefix:
                # Heading has extra trailing text — reset to canonical
                if paragraph.runs:
                    paragraph.runs[0].text = canonical
                    for run in paragraph.runs[1:]:
                        run.text = ""
                break


def _find_dafo_table(doc: Document):
    """Find the DAFO table by looking for keyword headers in its first 3 rows."""
    for table in doc.tables:
        for row in table.rows[:3]:
            for cell in row.cells:
                if any(kw in cell.text.strip().lower() for kw in _DAFO_HEADER_KEYWORDS):
                    return table
    return None


def _apply_dafo_direct(doc: Document, free_texts: dict[str, str]) -> None:
    """BUG 3 fix: Fill DAFO table cells by (row, col) position mapping."""
    dafo_table = _find_dafo_table(doc)
    if dafo_table is None:
        return
    rows = dafo_table.rows
    for (row_idx, col_idx), key in _DAFO_POSITIONS.items():
        value = free_texts.get(key, "")
        if row_idx < len(rows) and col_idx < len(rows[row_idx].cells):
            _set_cell_text(rows[row_idx].cells[col_idx], value)


def _apply_dpi_options(doc: Document, selections: dict[str, str | None]) -> None:
    """Clear all [Valorar] placeholders and clear col[1] placeholder text for option rows.

    Step 1: Clear every cell in doc.tables[1] whose text is exactly '[Valorar]'
            or '{{valorar}}' (the block/criterion header cells).
    Step 2: For each criterion in _OPTION_MAP, clear col[1] for all option rows
            (removes {{VALOR_CONSTITUCION}} etc.). No background changes — the
            template's pre-existing formatting is preserved via shutil.copy2.
    """
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    rows = table.rows
    n_rows = len(rows)

    # Step 1: Clear block/criterion header [Valorar] / {{valorar}} cells.
    _VALORAR_NORMS = {"[valorar]", "{{valorar}}"}
    for row in rows:
        seen: set[int] = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            if _normalize_text(cell.text) in _VALORAR_NORMS:
                _set_cell_text(cell, "")

    # Step 2: Per-criterion option row processing — write numeric score per option.
    for criterion, option_rows in _OPTION_MAP.items():
        # Get ordered option labels from scoring map (preserves matrix order)
        option_labels = list(SCORE_MAP.get(criterion, {}).keys())
        for j, row_idx in enumerate(option_rows):
            if row_idx >= n_rows:
                continue
            row_cells = rows[row_idx].cells
            if len(row_cells) < 2:
                continue

            if j < len(option_labels):
                score = SCORE_MAP[criterion][option_labels[j]]
                _set_cell_text(row_cells[1], str(score))
            else:
                # Non-scoring option row (recursos_economicos, personal_internacionalizacion)
                _set_cell_text(row_cells[1], "")


# ─── Public API ───────────────────────────────────────────────────────────────


def render_template(
    document_id: int,
    data: dict,
    empresa_name: str = "",
    document_type: str = "default",
) -> Path:
    """Fill the DPI template with direct fields, selections, and free texts.

    Args:
        document_id: Used to name the output file uniquely.
        data: dict with keys:
            - direct_fields: {Razon_Social, CIF, WEB, ...}
            - selections: {situacion_empresa: "opción elegida", ...}
            - free_texts: {definicion_potencial, conclusiones, dafo_*}
        empresa_name: Used in output filename (sanitized).
        document_type: Template variant name (default = plantilla.docx).

    Returns:
        Path to the generated DOCX file.
    """
    template_path = _find_template(document_type)

    # Build output filename — template original is NEVER touched
    empresa_clean = re.sub(r"[^\w\s-]", "", empresa_name).strip()
    empresa_clean = re.sub(r"\s+", " ", empresa_clean)
    filename = (
        f"Reporte {empresa_clean}.docx"
        if empresa_clean
        else f"Reporte {document_id}.docx"
    )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / filename

    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # BUG 1 fix: pass direct_fields as-is; _replace_in_paragraph handles case-insensitivity.
    # Add 'fecha' fallback (today's date) when the field is absent from the document.
    raw_direct = data.get("direct_fields", {})
    direct_fields: dict[str, str] = {k: (v or "") for k, v in raw_direct.items()}
    if not any(k.lower() == "fecha" for k in raw_direct):
        direct_fields["fecha"] = datetime.now().strftime("%d/%m/%Y")
    # CAMBIO 3: optional fields — ensure they map to "" so {{placeholders}} clear
    _OPTIONAL_FIELDS = (
        "CIF",
        "Cargo",
        "WEB",
        "Persona_Contacto",
        "Telefono_Contacto",
        "email",
        "VALOR_CONSTITUCION",
        "Reunion_Inicial",
        "Nombre_realizador",
        "sector",
        "producto_servicio",
    )
    for optional_key in _OPTIONAL_FIELDS:
        if optional_key not in direct_fields:
            direct_fields[optional_key] = ""

    # Clear DPI table header placeholders ({{valorar}}, {{VALOR_CONSTITUCION}}, etc.)
    for ph_key in _DPI_TABLE_PLACEHOLDERS:
        if not any(k.lower() == ph_key.lower() for k in direct_fields):
            direct_fields[ph_key] = ""

    selections: dict[str, str | None] = data.get("selections", {})
    free_texts: dict[str, str] = {
        k: (v or "") for k, v in data.get("free_texts", {}).items()
    }

    _clean_corrupted_headings(doc)  # reset headings with leaked text
    _apply_direct_fields_to_paragraphs(doc, direct_fields)
    _apply_dpi_options(doc, selections)  # write numeric scores for option rows
    _write_score_values(doc, selections)  # write selected score in golden cells
    _write_totals(doc, selections)  # block subtotals + DPI total
    _apply_dafo_direct(doc, free_texts)  # positional fill (BUG 3)
    _apply_free_texts_to_paragraphs(
        doc, free_texts
    )  # markers + {{dafo_*}} fallback (BUG 4)

    doc.save(output_path)
    return output_path
