import os
import json
from google import genai
from docx import Document

# --- 1. CONFIGURACIÓN ---
API_KEY = "AIzaSyDTH1P_-TXrYepmh6f6_t1faR_hqvlHeHU" # Pon tu llave aquí
MODEL_ID = "gemini-2.0-flash" 

client = genai.Client(api_key=API_KEY)

# Rutas para Windows (fíjate en la 'r' antes de las comillas para evitar errores de barras)
BASE_DIR = r"C:\Users\didix\Documents\Proyectos python\prueba 1"
RUTA_TXT = os.path.join(BASE_DIR, "reunion.txt")
RUTA_PLANTILLA = os.path.join(BASE_DIR, "Plantilla informe canarias - Copia.docx")
RUTA_SALIDA = os.path.join(BASE_DIR, "Informe_Final_Windows.docx")

# --- 2. FUNCIÓN PARA ANALIZAR CON IA ---
def analizar_datos(texto_reunion):
    print(f"🧠 La IA ({MODEL_ID}) está analizando la reunión...")
    
    # IMPORTANTE: Aquí definimos las etiquetas exactas que tienes en el Word
    prompt = f"""
    Basado en esta reunión: "{texto_reunion}"
    Extrae la información y devuélvela en formato JSON estrictamente.
    Las llaves del JSON deben ser EXACTAMENTE estas (respeta mayúsculas/minúsculas):
    "Razon_Social", "CIF", "WEB", "Persona_Contacto", "email", "fecha", "RESUMEN"
    """
    
    try:
        response = client.models.generate_content(model=MODEL_ID, contents=prompt)
        # Limpiamos el texto por si la IA añade bloques de código Markdown
        raw_text = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(raw_text)
    except Exception as e:
        print(f"❌ Error en la IA: {e}")
        return None

# --- 3. FUNCIÓN PARA RELLENAR EL WORD (Párrafos + Tablas) ---
def rellenar_word(datos_ia):
    if not datos_ia: return

    print("📝 Rellenando la plantilla Word...")
    doc = Document(RUTA_PLANTILLA)

    # Función interna para buscar y reemplazar
    def buscar_y_reemplazar(objeto_donde_buscar):
        for tag_simple, valor in datos_ia.items():
            etiqueta_completa = f"{{{{{tag_simple}}}}}" # Esto busca {{Razon_Social}}
            if etiqueta_completa in objeto_donde_buscar.text:
                print(f"✅ Reemplazando {etiqueta_completa}")
                objeto_donde_buscar.text = objeto_donde_buscar.text.replace(etiqueta_completa, str(valor))

    # BUSCAR EN PÁRRAFOS NORMALES
    for p in doc.paragraphs:
        buscar_y_reemplazar(p)

    # BUSCAR DENTRO DE LAS TABLAS (Tu CIF y Web suelen estar aquí)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    buscar_y_reemplazar(p)

    doc.save(RUTA_SALIDA)
    print(f"🎉 ¡HECHO! Archivo creado en: {RUTA_SALIDA}")

# --- 4. EJECUCIÓN PRINCIPAL ---
if __name__ == "__main__":
    if os.path.exists(RUTA_TXT) and os.path.exists(RUTA_PLANTILLA):
        with open(RUTA_TXT, "r", encoding="utf-8") as f:
            texto = f.read()
        
        resultado_ia = analizar_datos(texto)
        if resultado_ia:
            print(f"🔍 Datos encontrados por IA: {resultado_ia}")
            rellenar_word(resultado_ia)
    else:
        print("❌ Error: No se encuentra 'reunion.txt' o la plantilla en la carpeta.")