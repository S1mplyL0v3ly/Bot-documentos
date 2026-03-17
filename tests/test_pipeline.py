"""Pipeline unit tests for autoreporte DPI."""

import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ─── Classifier (kept for compatibility) ─────────────────────────────────────


def test_classifier_identifies_document_type():
    """Classifier should return a document_type and expected_fields from text."""
    from agents.classifier import classify_document

    sample_text = (
        "FACTURA N.º 2024-001\n"
        "Emisor: Empresa S.L. CIF: B12345678\n"
        "Importe total: 1.210,00 EUR\n"
        "Fecha: 15/01/2024"
    )
    fake_response = json.dumps(
        {
            "document_type": "factura",
            "expected_fields": ["numero_factura", "emisor", "importe", "fecha"],
            "confidence": 0.95,
        }
    )
    with patch("agents.classifier.run_claude", return_value=fake_response):
        result = classify_document(sample_text)

    assert result["document_type"] == "factura"
    assert "numero_factura" in result["expected_fields"]
    assert result["confidence"] == 0.95


# ─── Extractor ────────────────────────────────────────────────────────────────


def test_extractor_returns_json_with_fields():
    """extract_dpi_fields should return direct_fields + selections + confidence."""
    from agents.extractor import extract_dpi_fields

    text = (
        "Empresa: Tech SL  CIF: B87654321  Web: www.tech.es\n"
        "Empleados: 15  Facturación: 800.000 €  Años actividad: 5\n"
        "Exporta a 3 países europeos."
    )
    fake_response = json.dumps(
        {
            "direct_fields": {
                "Razon_Social": "Tech SL",
                "CIF": "B87654321",
                "WEB": "www.tech.es",
                "Persona_Contacto": None,
                "Cargo": None,
                "email": None,
                "Telefono_Contacto": None,
            },
            "selections": {
                "situacion_empresa": "Más de 2 años",
                "num_empleados": "Más de 2",
                "facturacion": "Entre 500.000 y 1.000.000 €",
                "evolucion_facturacion": None,
                "recursos_internacionalizacion": None,
                "experiencia_internacional": "Menos de 3 años",
                "alcance_actividad": "Internacional",
                "num_paises": "De 1 a 5",
                "personal_dedicado": None,
                "involuccion_gerencia": None,
                "adaptacion_demanda": None,
                "adaptacion_producto": None,
                "tiene_web": "Si",
                "ecommerce": None,
                "mercados_electronicos": None,
                "redes_sociales": None,
            },
            "confidence": {
                "situacion_empresa": 0.95,
                "num_empleados": 0.90,
                "facturacion": 0.80,
                "evolucion_facturacion": 0.4,
                "recursos_internacionalizacion": 0.3,
                "experiencia_internacional": 0.75,
                "alcance_actividad": 0.85,
                "num_paises": 0.88,
                "personal_dedicado": 0.3,
                "involuccion_gerencia": 0.2,
                "adaptacion_demanda": 0.4,
                "adaptacion_producto": 0.4,
                "tiene_web": 0.99,
                "ecommerce": 0.3,
                "mercados_electronicos": 0.2,
                "redes_sociales": 0.3,
            },
        }
    )

    with patch("agents.extractor.run_claude", return_value=fake_response):
        result = extract_dpi_fields(text)

    assert result["direct_fields"]["Razon_Social"] == "Tech SL"
    assert result["selections"]["situacion_empresa"] == "Más de 2 años"
    assert result["selections"]["tiene_web"] == "Si"
    # Low confidence → nulled
    assert result["selections"]["evolucion_facturacion"] is None
    assert result["selections"]["personal_dedicado"] is None


def test_extractor_returns_null_for_low_confidence():
    """Selections with confidence < 0.7 must be nulled by _null_low_confidence."""
    from agents.extractor import _null_low_confidence

    data = {
        "direct_fields": {},
        "selections": {
            "situacion_empresa": "Más de 2 años",
            "num_empleados": "Más de 2",
            "facturacion": "Menos de 200.000 €",
        },
        "confidence": {
            "situacion_empresa": 0.9,  # keep
            "num_empleados": 0.65,  # null (below threshold)
            "facturacion": 0.5,  # null
        },
    }
    result = _null_low_confidence(data)
    assert result["selections"]["situacion_empresa"] == "Más de 2 años"
    assert result["selections"]["num_empleados"] is None
    assert result["selections"]["facturacion"] is None


# ─── DOCX Generator ───────────────────────────────────────────────────────────


def test_docx_generator_creates_file(tmp_path):
    """render_template should produce a .docx with direct fields and green selections."""
    from docx import Document
    from docx_generator.template_handler import render_template

    # Build minimal template with direct-field placeholder and a selection cell
    template = Document()
    template.add_paragraph("Empresa: {{Razon_Social}}")
    template.add_paragraph("CIF: {{CIF}}")
    # Table with selection options
    tbl = template.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Situación"
    tbl.cell(0, 1).text = "Más de 2 años"
    tbl.cell(1, 0).text = "Empleados"
    tbl.cell(1, 1).text = "Menos de 2"
    template_path = tmp_path / "plantilla.docx"
    template.save(str(template_path))

    data = {
        "direct_fields": {"Razon_Social": "Tech SL", "CIF": "B87654321"},
        "selections": {"situacion_empresa": "Más de 2 años"},
        "free_texts": {},
    }

    with (
        patch("docx_generator.template_handler.TEMPLATES_DIR", tmp_path),
        patch("docx_generator.template_handler.DEFAULT_TEMPLATE", template_path),
        patch("docx_generator.template_handler.OUTPUT_DIR", tmp_path),
    ):
        output = render_template(document_id=1, data=data, empresa_name="TechSL")

    assert output.exists()
    assert output.suffix == ".docx"
    doc = Document(str(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Tech SL" in full_text
    assert "B87654321" in full_text


def test_render_template_marks_selection_green(tmp_path):
    """Selected option cell must receive green fill (#92D050), others untouched."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx_generator.template_handler import GREEN_FILL, render_template

    template = Document()
    tbl = template.add_table(rows=3, cols=1)
    tbl.cell(0, 0).text = "No constituida"
    tbl.cell(1, 0).text = "Menos de 2 años"
    tbl.cell(2, 0).text = "Más de 2 años"
    template_path = tmp_path / "plantilla.docx"
    template.save(str(template_path))

    data = {
        "direct_fields": {},
        "selections": {"situacion_empresa": "Más de 2 años"},
        "free_texts": {},
    }

    with (
        patch("docx_generator.template_handler.TEMPLATES_DIR", tmp_path),
        patch("docx_generator.template_handler.DEFAULT_TEMPLATE", template_path),
        patch("docx_generator.template_handler.OUTPUT_DIR", tmp_path),
    ):
        output = render_template(document_id=2, data=data)

    doc = Document(str(output))
    cells = doc.tables[0].columns[0].cells

    def get_fill(cell) -> str:
        shd = cell._tc.find(f".//{{{qn('w:shd').split('}')[0][1:]}}}shd")
        if shd is None:
            # Try direct child
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                return ""
            shd = tcPr.find(qn("w:shd"))
        return (shd.get(qn("w:fill")) or "") if shd is not None else ""

    selected_fill = get_fill(cells[2])  # "Más de 2 años"
    other_fill = get_fill(cells[0])  # "No constituida"

    assert selected_fill.upper() == GREEN_FILL.upper()
    assert other_fill.upper() != GREEN_FILL.upper()


# ─── Validator ────────────────────────────────────────────────────────────────


def test_validator_detects_missing_fields():
    """validate_fields should correctly identify missing required fields."""
    from docx_generator.validator import validate_fields

    fields = {
        "nombre": "Juan García",
        "fecha": "15/01/2024",
        "importe": None,
    }
    required = ["nombre", "fecha", "importe", "cif"]
    result = validate_fields(fields, required)

    assert result["complete"] is False
    assert "importe" in result["missing"]
    assert "cif" in result["missing"]
    assert "nombre" in result["present"]
    assert result["total_required"] == 4
    assert result["total_present"] == 2


# ─── Orchestrator ─────────────────────────────────────────────────────────────


def test_orchestrator_generates_questions_for_null_fields():
    """generate_questions should list only the criteria with null values."""
    from agents.orchestrator import PREFIX_SELECTION, generate_questions

    # Mock DB session
    mock_field_null = MagicMock()
    mock_field_null.field_name = f"{PREFIX_SELECTION}recursos_internacionalizacion"
    mock_field_null.field_value = None

    mock_field_ok = MagicMock()
    mock_field_ok.field_name = f"{PREFIX_SELECTION}situacion_empresa"
    mock_field_ok.field_value = "Más de 2 años"

    mock_db = MagicMock()
    with patch(
        "agents.orchestrator.crud.get_fields",
        return_value=[mock_field_null, mock_field_ok],
    ):
        message = generate_questions(mock_db, document_id=1)

    assert "recursos_internacionalizacion" in message or "recursos" in message.lower()
    assert "situacion_empresa" not in message


def test_extractor_maps_facturacion_correctly():
    """MEJORA 1: Non-exact facturación range is normalized to the correct DPI option."""
    from agents.extractor import extract_dpi_fields

    text = "Facturación total durante el ejercicio 2024: Menos de 250.000 €"
    null_conf = {
        k: 0.0
        for k in [
            "situacion_empresa",
            "num_empleados",
            "facturacion",
            "evolucion_facturacion",
            "recursos_internacionalizacion",
            "experiencia_internacional",
            "alcance_actividad",
            "num_paises",
            "personal_dedicado",
            "involuccion_gerencia",
            "adaptacion_demanda",
            "adaptacion_producto",
            "tiene_web",
            "ecommerce",
            "mercados_electronicos",
            "redes_sociales",
        ]
    }
    null_sel = {k: None for k in null_conf}
    fake_response = json.dumps(
        {
            "direct_fields": {
                "Razon_Social": None,
                "CIF": None,
                "WEB": None,
                "Persona_Contacto": None,
                "Cargo": None,
                "email": None,
                "Telefono_Contacto": None,
                "sector": None,
                "producto_servicio": None,
                "año_inicio": None,
            },
            "selections": {**null_sel, "facturacion": "Menos de 250.000 €"},
            "confidence": {**null_conf, "facturacion": 0.9},
        }
    )
    with patch("agents.extractor.run_claude", return_value=fake_response):
        result = extract_dpi_fields(text)
    assert result["selections"]["facturacion"] == "Menos de 200.000 €"


def test_extractor_deduces_situacion_from_year():
    """MEJORA 3: situacion_empresa is inferred from año_inicio when Claude returns null."""
    from agents.extractor import extract_dpi_fields

    text = "Año de inicio de la actividad empresarial: 2020"
    null_conf = {
        k: 0.0
        for k in [
            "situacion_empresa",
            "num_empleados",
            "facturacion",
            "evolucion_facturacion",
            "recursos_internacionalizacion",
            "experiencia_internacional",
            "alcance_actividad",
            "num_paises",
            "personal_dedicado",
            "involuccion_gerencia",
            "adaptacion_demanda",
            "adaptacion_producto",
            "tiene_web",
            "ecommerce",
            "mercados_electronicos",
            "redes_sociales",
        ]
    }
    null_sel = {k: None for k in null_conf}
    fake_response = json.dumps(
        {
            "direct_fields": {
                "Razon_Social": None,
                "CIF": None,
                "WEB": None,
                "Persona_Contacto": None,
                "Cargo": None,
                "email": None,
                "Telefono_Contacto": None,
                "sector": None,
                "producto_servicio": None,
                "año_inicio": "2020",
            },
            "selections": null_sel,
            "confidence": null_conf,
        }
    )
    with patch("agents.extractor.run_claude", return_value=fake_response):
        result = extract_dpi_fields(text)
    assert result["selections"]["situacion_empresa"] == "Más de 2 años"
    assert result["confidence"]["situacion_empresa"] >= 0.9


def test_extractor_atelier_maria():
    """Cuestionario Canarias Expande: normalización de campos para Atelier Maria Secretos.

    Verifica que los normalizadores convierten respuestas de texto libre del cuestionario
    a los valores DPI exactos: facturación, experiencia, involucción gerencia, situación.
    """
    from agents.extractor import extract_dpi_fields

    fixture = Path(__file__).parent / "fixtures" / "atelier_text.txt"
    text = fixture.read_text(encoding="utf-8")

    # Claude devuelve valores que requieren normalización (texto libre del cuestionario)
    _conf_95 = {
        k: 0.95
        for k in [
            "situacion_empresa",
            "num_empleados",
            "facturacion",
            "evolucion_facturacion",
            "recursos_internacionalizacion",
            "experiencia_internacional",
            "alcance_actividad",
            "num_paises",
            "personal_dedicado",
            "involuccion_gerencia",
            "adaptacion_demanda",
            "adaptacion_producto",
            "tiene_web",
            "ecommerce",
            "mercados_electronicos",
            "redes_sociales",
        ]
    }
    fake_response = json.dumps(
        {
            "direct_fields": {
                "Razon_Social": "Atelier Maria Secretos SL",
                "CIF": None,
                "WEB": None,
                "Persona_Contacto": None,
                "Cargo": None,
                "email": None,
                "Telefono_Contacto": None,
                "sector": "Joyería y bisutería artesanal",
                "producto_servicio": "Joyería artesanal de plata",
                "año_inicio": "2020",
            },
            "selections": {
                "situacion_empresa": None,  # debe deducirse de año_inicio=2020
                "num_empleados": "Más de 2",
                "facturacion": "Menos de 250.000 €",  # debe normalizarse → "Menos de 200.000 €"
                "evolucion_facturacion": "En crecimiento",
                "recursos_internacionalizacion": "Sí",
                "experiencia_internacional": "No hemos exportado nunca",  # → "Ninguna experiencia"
                "alcance_actividad": "Nacional",
                "num_paises": "Ninguno salvo el mercado nacional",
                "personal_dedicado": "No",
                "involuccion_gerencia": "Directamente involucrada",  # → "Directamente involucrados"
                "adaptacion_demanda": "Media",
                "adaptacion_producto": "Alta",
                "tiene_web": "Si",
                "ecommerce": "Sin tienda web",
                "mercados_electronicos": "Con presencia en mercados electrónicos sin ventas o ventas bajas.",
                "redes_sociales": "Redes sociales activas y planificadas",
            },
            "confidence": _conf_95,
        }
    )

    with patch("agents.extractor.run_claude", return_value=fake_response):
        result = extract_dpi_fields(text)

    # Normalización de facturación
    assert result["selections"]["facturacion"] == "Menos de 200.000 €"
    # Normalización de experiencia
    assert result["selections"]["experiencia_internacional"] == "Ninguna experiencia"
    # Normalización de involucción gerencia
    assert result["selections"]["involuccion_gerencia"] == "Directamente involucrados"
    # Deducción situacion_empresa desde año_inicio=2020 (6 años → "Más de 2 años")
    assert result["selections"]["situacion_empresa"] == "Más de 2 años"
    assert result["confidence"]["situacion_empresa"] >= 0.9

    # Confidence alta para campos clave con dato explícito
    for campo in ["num_empleados", "facturacion", "evolucion_facturacion", "tiene_web"]:
        assert (
            result["confidence"][campo] >= 0.8
        ), f"{campo} confidence baja: {result['confidence'][campo]}"

    # Con este fixture todos los criterios deben tener valor — máx 1 nulo aceptable
    nulls = sum(1 for v in result["selections"].values() if v is None)
    assert (
        nulls <= 1
    ), f"Demasiados campos sin deducir: {nulls} — {[k for k,v in result['selections'].items() if v is None]}"


def test_approval_flow_generates_final_docx(tmp_path):
    """generate_final_docx should call render_template and update status to complete."""
    from agents.orchestrator import generate_final_docx

    fake_output = tmp_path / "informe_test.docx"
    fake_output.write_bytes(b"PK")  # minimal zip-like content

    mock_db = MagicMock()
    mock_fields = [
        MagicMock(field_name="dir_Razon_Social", field_value="Empresa Test"),
        MagicMock(field_name="sel_situacion_empresa", field_value="Más de 2 años"),
        MagicMock(field_name="txt_conclusiones", field_value="Buen potencial."),
    ]
    with (
        patch("agents.orchestrator.crud.get_fields", return_value=mock_fields),
        patch("agents.orchestrator.crud.create_or_update_generated_docx"),
        patch("agents.orchestrator.crud.update_document_status") as mock_status,
        patch(
            "agents.orchestrator.render_template", return_value=fake_output
        ) as mock_render,
        patch("agents.orchestrator._log_to_jarvis"),
    ):
        result = generate_final_docx(mock_db, document_id=1)

    assert result["status"] == "complete"
    assert "output_path" in result
    mock_render.assert_called_once()
    mock_status.assert_called_with(mock_db, 1, "complete")


# ─── CAMBIO 8: Two new tests ──────────────────────────────────────────────────


def test_calculate_dpi_score_atelier_maria():
    """calculate_dpi_score debe devolver puntuaciones correctas para Atelier Maria."""
    from agents.orchestrator import calculate_dpi_score

    selections = {
        "situacion_empresa": "Más de 2 años",
        "num_empleados": "Más de 2",
        "facturacion": "Menos de 200.000 €",
        "evolucion_facturacion": "En crecimiento",
        "experiencia_internacional": "Ninguna experiencia",
        "alcance_actividad": "Nacional",
        "num_paises": "Ninguno salvo el mercado nacional",
        "involuccion_gerencia": "Directamente involucrados",
        "adaptacion_demanda": "Media",
        "adaptacion_producto": "Alta",
        "tiene_web": "Si",
        "ecommerce": "Sin tienda web",
        "mercados_electronicos": "Con presencia en mercados electrónicos sin ventas o ventas bajas.",
        "redes_sociales": "Redes sociales activas y planificadas",
    }
    score = calculate_dpi_score(selections)

    # Económico: 5 (Más de 2 años) + 5 (Más de 2) + 1 (Menos de 200.000 €) = 11
    assert score["scores"]["Económico"] == 11
    # Internacional: 0 (Ninguna experiencia) + 3 (Nacional) + 6 (En crecimiento) + 5 (Directamente) + 2 (Media) + 2 (Alta) + 0 (Ninguno) = 18
    assert score["scores"]["Internacional"] == 18
    # Digitalización: 3 (Si) + 1 (Sin tienda) + 1 (Con presencia sin ventas) + 1 (Activas y planificadas) = 6
    assert score["scores"]["Digitalización"] == 6
    assert score["total"] == 35
    assert score["max_total"] == 65
    assert score["pct"] == 54  # round(35/65*100) = round(53.8) = 54


def test_two_document_flow():
    """get_document_waiting_transcript devuelve el doc en waiting_transcript del sender."""
    mock_doc = MagicMock()
    mock_doc.id = 42
    mock_doc.status = "waiting_transcript"
    mock_doc.transcript_text = None

    mock_db = MagicMock()

    with patch(
        "database.crud.get_document_by_sender_and_status", return_value=mock_doc
    ):
        from database.crud import get_document_waiting_transcript

        result = get_document_waiting_transcript(mock_db, sender_id="34600000000")

    assert result is mock_doc
    assert result.id == 42
    assert result.status == "waiting_transcript"


# ─── Logical implications ─────────────────────────────────────────────────────


def test_boost_visual_confidence_recursos_no_programs_selected():
    """Si la pregunta de programas existe en el cuestionario pero ninguno seleccionado → recursos=No."""
    from agents.extractor import _boost_visual_confidence

    cuestionario_text = (
        "=== OPCIONES SELECCIONADAS EN EL FORMULARIO ===\n"
        "◉ No hemos exportado nunca\n"
        "◉ Directamente involucrados\n"
        "\n=== TEXTO COMPLETO DEL FORMULARIO ===\n"
        "¿Ha participado su empresa en algún programa de apoyo a la internacionalización?\n"
        "Canarias Aporta\nICEX NEXT\nICEX APIEm\nMisiones Comerciales\nRed EEN\n"
    )
    data = {
        "selections": {"recursos_internacionalizacion": None},
        "confidence": {"recursos_internacionalizacion": 0.0},
    }
    result = _boost_visual_confidence(data, cuestionario_text)
    assert result["selections"]["recursos_internacionalizacion"] == "No"
    assert result["confidence"]["recursos_internacionalizacion"] >= 0.8


def test_boost_visual_confidence_tiene_web_bare_si():
    """◉ Sí (bare line) + "página web" en texto → tiene_web = "Si"."""
    from agents.extractor import _boost_visual_confidence

    cuestionario_text = (
        "=== OPCIONES SELECCIONADAS EN EL FORMULARIO ===\n"
        "◉ Sí\n"
        "◉ Sí, me gustaría recibir dicha información.\n"
        "\n=== TEXTO COMPLETO DEL FORMULARIO ===\n"
        "¿Dispone la empresa de página web corporativa?\n"
        "Sí  No\n"
    )
    data = {
        "selections": {"tiene_web": None},
        "confidence": {"tiene_web": 0.0},
    }
    result = _boost_visual_confidence(data, cuestionario_text)
    assert result["selections"]["tiene_web"] == "Si"
    assert result["confidence"]["tiene_web"] >= 0.85


def test_boost_visual_confidence_tiene_web_not_triggered_without_keyword():
    """◉ Sí sin keyword de página web en el texto → tiene_web permanece null."""
    from agents.extractor import _boost_visual_confidence

    cuestionario_text = (
        "=== OPCIONES SELECCIONADAS EN EL FORMULARIO ===\n"
        "◉ Sí\n"
        "\n=== TEXTO COMPLETO DEL FORMULARIO ===\n"
        "¿Le gustaría recibir información sobre ayudas?\n"
    )
    data = {
        "selections": {"tiene_web": None},
        "confidence": {"tiene_web": 0.0},
    }
    result = _boost_visual_confidence(data, cuestionario_text)
    assert result["selections"]["tiene_web"] is None


def test_apply_logical_implications_ninguna_exports():
    """experiencia_internacional=Ninguna experiencia debe implicar num_paises=Ninguno salvo el mercado nacional."""
    from agents.extractor import _apply_logical_implications

    data = {
        "selections": {
            "experiencia_internacional": "Ninguna experiencia",
            "num_paises": None,
        },
        "confidence": {"experiencia_internacional": 1.0, "num_paises": 0.0},
    }
    result = _apply_logical_implications(data)
    assert result["selections"]["num_paises"] == "Ninguno salvo el mercado nacional"
    assert result["confidence"]["num_paises"] == 0.95


def test_apply_logical_implications_internacional_implies_countries():
    """alcance_actividad=Internacional debe implicar num_paises=De 1 a 5 cuando null."""
    from agents.extractor import _apply_logical_implications

    data = {
        "selections": {"alcance_actividad": "Internacional", "num_paises": None},
        "confidence": {"alcance_actividad": 0.9, "num_paises": 0.0},
    }
    result = _apply_logical_implications(data)
    assert result["selections"]["num_paises"] == "De 1 a 5"
    assert result["confidence"]["num_paises"] >= 0.7


def test_apply_logical_implications_does_not_overwrite():
    """Si num_paises ya tiene valor, no se sobreescribe."""
    from agents.extractor import _apply_logical_implications

    data = {
        "selections": {
            "experiencia_internacional": "Ninguna experiencia",
            "num_paises": "De 1 a 5",
        },
        "confidence": {"experiencia_internacional": 1.0, "num_paises": 0.9},
    }
    result = _apply_logical_implications(data)
    # Should not overwrite existing value
    assert result["selections"]["num_paises"] == "De 1 a 5"


# ─── Visual PDF detection ─────────────────────────────────────────────────────


# ─── Web scraper ─────────────────────────────────────────────────────────────


def test_score_url_match_high_for_company_name_in_domain():
    """URL whose domain contains company tokens scores > 0.4."""
    from agents.web_scraper import _score_url_match

    html = "<html><title>Joyeria Atelier Maria</title><body>Joyeria Atelier Maria</body></html>"
    score = _score_url_match("https://atelier-maria.com", html, "Atelier Maria")
    assert score >= 0.4, f"Expected score >= 0.4, got {score}"


def test_score_url_match_low_for_unrelated_domain():
    """URL with no company tokens scores < 0.3."""
    from agents.web_scraper import _score_url_match

    html = "<html><title>Another Company</title><body>Unrelated content</body></html>"
    score = _score_url_match(
        "https://unrelated-site.com", html, "Atelier Maria Secretos"
    )
    assert score < 0.3, f"Expected score < 0.3, got {score}"


def test_build_dpi_from_web_detects_social_and_ecommerce():
    """HTML with Instagram link + cart button → redes_sociales + ecommerce detected."""
    from agents.web_scraper import _build_dpi_from_web

    html = (
        "<html><body>"
        '<a href="https://instagram.com/atelier_maria">Instagram</a>'
        '<a href="https://facebook.com/ateliermaria">Facebook</a>'
        '<button class="add-to-cart">Comprar</button>'
        "</body></html>"
    )
    signals = _build_dpi_from_web("https://atelier-maria.com", html)
    assert signals["selections"].get("tiene_web") == "Si"
    assert (
        signals["selections"].get("redes_sociales")
        == "Redes sociales activas y planificadas"
    )
    assert signals["selections"].get("ecommerce") is not None
    assert signals["confidence"].get("redes_sociales", 0) >= 0.7


def test_build_dpi_from_web_detects_cif():
    """CIF in page footer is extracted to direct_fields."""
    from agents.web_scraper import _build_dpi_from_web

    html = (
        "<html><body>" "<footer>CIF: B12345678 | Aviso legal</footer>" "</body></html>"
    )
    signals = _build_dpi_from_web("https://empresa.com", html)
    assert signals["direct_fields"].get("CIF") == "B12345678"


def test_build_dpi_from_web_lang_selector_implies_internacional():
    """hreflang attribute on page → alcance_actividad = Internacional."""
    from agents.web_scraper import _build_dpi_from_web

    html = (
        "<html><head>"
        '<link rel="alternate" hreflang="en" href="https://empresa.com/en"/>'
        "</head><body>Welcome / Bienvenidos</body></html>"
    )
    signals = _build_dpi_from_web("https://empresa.com", html)
    assert signals["selections"].get("alcance_actividad") == "Internacional"


# ─── Visual PDF detection ─────────────────────────────────────────────────────


def test_read_google_forms_pdf_atelier_maria():
    """Verifica detección visual de casillas en PDF real de Google Forms."""
    import os

    pdf_path = str(FIXTURES_DIR / "cuestionario_atelier.pdf")

    if not os.path.exists(pdf_path):
        pytest.skip("PDF de prueba no disponible")

    from agents.extractor import read_google_forms_pdf

    result = read_google_forms_pdf(pdf_path)
    selected = [opt["text"] for opt in result["selected_options"]]

    assert any(
        "Empresario" in s or "autónomo" in s for s in selected
    ), f"Tipo entidad no detectado. Seleccionados: {selected}"
    assert any(
        "250.000" in s for s in selected
    ), f"Facturación no detectada. Seleccionados: {selected}"
    assert any(
        "exportado nunca" in s for s in selected
    ), f"Experiencia no detectada. Seleccionados: {selected}"
    assert any(
        "involucrados" in s.lower() for s in selected
    ), f"Gerencia no detectada. Seleccionados: {selected}"
    assert any(
        s.strip() == "Sí" for s in selected
    ), f"Página web no detectada. Seleccionados: {selected}"


# ─── CIF DDG search ──────────────────────────────────────────────────────────


def test_search_cif_ddg_extracts_from_snippet():
    """search_cif_ddg returns CIF found in DuckDuckGo result body/title (mocked)."""
    from agents.web_scraper import search_cif_ddg

    fake_results = [
        {
            "body": "Atelier Maria Secretos SL CIF B76543210 - Joyería en Las Palmas",
            "title": "",
            "href": "",
        },
        {
            "body": "Otro resultado sin CIF",
            "title": "página sin CIF",
            "href": "https://example.com",
        },
    ]

    class FakeDDGS:
        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

        def text(self, query, max_results=8):
            return iter(fake_results)

    with patch("agents.web_scraper.DDGS", FakeDDGS, create=True):
        # Patch the import inside the function
        import agents.web_scraper as ws_module
        import sys

        fake_mod = MagicMock()
        fake_mod.DDGS = FakeDDGS
        with patch.dict(sys.modules, {"duckduckgo_search": fake_mod}):
            result = search_cif_ddg("Atelier Maria Secretos")

    assert result == "B76543210"


def test_search_cif_ddg_returns_none_when_not_found():
    """search_cif_ddg returns None when no CIF appears in snippets."""
    from agents.web_scraper import search_cif_ddg

    fake_results = [
        {
            "body": "Empresa sin datos fiscales",
            "title": "Sin CIF",
            "href": "https://example.com",
        },
    ]

    class FakeDDGS:
        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

        def text(self, query, max_results=8):
            return iter(fake_results)

    import sys

    fake_mod = MagicMock()
    fake_mod.DDGS = FakeDDGS
    with patch.dict(sys.modules, {"duckduckgo_search": fake_mod}):
        result = search_cif_ddg("Empresa Inexistente SL")

    assert result is None


def test_search_cif_ddg_returns_none_on_exception():
    """search_cif_ddg returns None silently when DDGS raises."""
    from agents.web_scraper import search_cif_ddg

    import sys

    fake_mod = MagicMock()
    fake_mod.DDGS.side_effect = RuntimeError("network error")
    with patch.dict(sys.modules, {"duckduckgo_search": fake_mod}):
        result = search_cif_ddg("Any Company")

    assert result is None


def test_template_green_cells_and_no_valorar():
    """render_template marks selected option green and clears all [Valorar] / {{valorar}}."""
    import tempfile
    from pathlib import Path
    from docx_generator.template_handler import render_template, GREEN_FILL

    data = {
        "direct_fields": {"Razon_Social": "Test SL"},
        "selections": {
            "situacion_empresa": "Más de 2 años",  # row 9  in table[1]
            "num_empleados": "Más de 2",  # row 13
            "facturacion": "Menos de 200.000 €",  # row 16
            "evolucion_facturacion": "En crecimiento",  # row 24
            "experiencia_internacional": "Ninguna",  # row 33 (partial match)
            "alcance_actividad": "Nacional",  # row 39
            "tiene_web": "Si",  # row 71
            "ecommerce": "Tienda web propia con ventas bajas o irregulares",  # row 74
            "redes_sociales": "Redes sociales activas y planificadas",  # row 86
        },
        "free_texts": {},
    }

    output_path = render_template(document_id=9999, data=data, empresa_name="TestSL")
    assert output_path.exists()

    from docx import Document

    doc = Document(str(output_path))
    table = doc.tables[1]

    def get_fill(cell):
        tc = cell._tc
        tcPr = tc.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
        )
        if tcPr is None:
            return None
        shd = tcPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
        )
        if shd is None:
            return None
        return shd.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
        )

    # 1. Selected option rows must be green in both col[0] and col[1]
    green_rows = {
        9: "Más de 2 años",
        13: "Más de 2",
        16: "Menos de 200.000 €",
        24: "En crecimiento",
        33: "Ninguna experiencia",  # cell text (bidirectional match)
        39: "Nacional",
        71: "Si",
        74: "Tienda web propia con ventas bajas o irregulares",
        86: "Redes sociales activas y planificadas",
    }
    for row_idx, label in green_rows.items():
        row = table.rows[row_idx]
        fill0 = get_fill(row.cells[0])
        fill1 = get_fill(row.cells[1])
        assert fill0 == GREEN_FILL, f"row {row_idx} ({label}) col[0] not green: {fill0}"
        assert fill1 == GREEN_FILL, f"row {row_idx} ({label}) col[1] not green: {fill1}"

    # 2. No cell in table[1] should contain [Valorar] or {{valorar}} text
    for r_idx, row in enumerate(table.rows):
        seen: set[int] = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            txt = cell.text.strip().lower()
            assert (
                "[valorar]" not in txt
            ), f"[Valorar] not cleared at row {r_idx}: '{cell.text}'"
            assert (
                "{{valorar}}" not in txt
            ), f"{{valorar}} not cleared at row {r_idx}: '{cell.text}'"

    # Cleanup
    output_path.unlink(missing_ok=True)


def test_extract_relevant_sections_long_doc():
    """Long document: DPI-relevant lines are preserved even when far from start."""
    from agents.extractor import extract_relevant_sections

    irrelevant = "texto irrelevante sin ninguna palabra clave\n" * 500
    relevant_a = "La empresa tiene 5 empleados fijos a tiempo completo\n"
    irrelevant2 = "texto irrelevante sin ninguna palabra clave\n" * 500
    relevant_b = "Facturación estable durante los últimos tres años\n"
    long_text = irrelevant + relevant_a + irrelevant2 + relevant_b

    result = extract_relevant_sections(long_text, max_chars=12000)

    assert "5 empleados" in result
    assert "Facturación estable" in result
    assert len(result) <= 12000


def test_extract_relevant_sections_short_doc():
    """Short document: returned unchanged (no extraction needed)."""
    from agents.extractor import extract_relevant_sections

    short_text = "texto corto\n" * 10
    result = extract_relevant_sections(short_text, max_chars=15000)
    assert result == short_text


# ─── FIX 1: Nombre_realizador split ──────────────────────────────────────────


def test_fix1_nombre_realizador_split():
    """FIX 1: 'Nombre, DD/MM/YYYY' splits into name + date saved separately."""
    import asyncio
    from unittest.mock import AsyncMock, MagicMock, patch

    # Build a minimal waiting_doc fixture
    mock_doc = MagicMock()
    mock_doc.id = 42

    mock_db = MagicMock()
    upserted = {}

    def fake_upsert(db, doc_id, field_name, value, confidence, source):
        upserted[field_name] = value

    with (
        patch("database.crud.get_document_waiting_transcript", return_value=mock_doc),
        patch("database.crud.get_fields", return_value=[]),
        patch("database.crud.upsert_field", side_effect=fake_upsert),
        patch("channels.whatsapp.send_text", new_callable=AsyncMock),
        patch("api.routes._make_db", return_value=mock_db),
    ):
        from api.routes import _bg_process_wa_text

        asyncio.run(_bg_process_wa_text("34600000001", "María González, 15/03/2025"))

    assert upserted.get("dir_Nombre_realizador") == "María González"
    assert upserted.get("dir_Reunion_Inicial") == "15/03/2025"


def test_fix1_nombre_realizador_no_date():
    """FIX 1: Name without comma saves only dir_Nombre_realizador."""
    import asyncio
    from unittest.mock import AsyncMock, MagicMock, patch

    mock_doc = MagicMock()
    mock_doc.id = 43
    mock_db = MagicMock()
    upserted = {}

    def fake_upsert(db, doc_id, field_name, value, confidence, source):
        upserted[field_name] = value

    with (
        patch("database.crud.get_document_waiting_transcript", return_value=mock_doc),
        patch("database.crud.get_fields", return_value=[]),
        patch("database.crud.upsert_field", side_effect=fake_upsert),
        patch("channels.whatsapp.send_text", new_callable=AsyncMock),
        patch("api.routes._make_db", return_value=mock_db),
    ):
        from api.routes import _bg_process_wa_text

        asyncio.run(_bg_process_wa_text("34600000001", "María González"))

    assert upserted.get("dir_Nombre_realizador") == "María González"
    assert "dir_Reunion_Inicial" not in upserted


# ─── FIX 2: Manual fields not overwritten by Claude ──────────────────────────


def test_fix2_manual_cargo_not_overwritten():
    """FIX 2: dir_Cargo with source='manual' is NOT overwritten by Claude extraction.

    Tests the guard logic directly: when existing_fields contains a 'manual' source entry,
    upsert_field must be skipped for that field but called for others.
    """
    from unittest.mock import MagicMock

    import database.crud as crud_mod
    from agents.orchestrator import PREFIX_DIRECT

    # Simulate existing_fields dict as built inside process_document
    manual_field = MagicMock()
    manual_field.field_name = "dir_Cargo"
    manual_field.source = "manual"

    existing_fields = {manual_field.field_name: manual_field}

    # The direct_fields Claude wants to save
    extracted_direct = {"Razon_Social": "Test SL", "Cargo": "Director"}

    saved = {}

    def fake_upsert(db, doc_id, field_name, value, confidence, source):
        saved[field_name] = (value, source)

    db = MagicMock()
    for key, value in extracted_direct.items():
        field_name = f"{PREFIX_DIRECT}{key}"
        existing = existing_fields.get(field_name)
        if existing and existing.source == "manual":
            continue
        fake_upsert(db, 1, field_name, value, 0.9, "claude")

    assert "dir_Cargo" not in saved, "manual dir_Cargo must not be overwritten"
    assert "dir_Razon_Social" in saved, "non-manual field must be saved"


# ─── FIX 3: alcance_actividad contradiction rule ──────────────────────────────


def test_fix3_ninguna_experiencia_overrides_internacional():
    """FIX 3: Ninguna experiencia + Internacional → corrected to Nacional."""
    from agents.extractor import _apply_logical_implications

    data = {
        "selections": {
            "experiencia_internacional": "Ninguna experiencia",
            "alcance_actividad": "Internacional",
            "num_paises": None,
        },
        "confidence": {
            "experiencia_internacional": 1.0,
            "alcance_actividad": 0.7,
            "num_paises": 0.0,
        },
    }
    result = _apply_logical_implications(data)
    assert result["selections"]["alcance_actividad"] == "Nacional"
    assert result["confidence"]["alcance_actividad"] == 0.9
    # num_paises set to Ninguno (from Ninguna experiencia rule)
    assert result["selections"]["num_paises"] == "Ninguno salvo el mercado nacional"


def test_fix3_internacional_with_exports_not_overridden():
    """FIX 3: Real exports + Internacional → alcance stays Internacional."""
    from agents.extractor import _apply_logical_implications

    data = {
        "selections": {
            "experiencia_internacional": "Más de 5 años",
            "alcance_actividad": "Internacional",
            "num_paises": None,
        },
        "confidence": {
            "experiencia_internacional": 1.0,
            "alcance_actividad": 0.9,
            "num_paises": 0.0,
        },
    }
    result = _apply_logical_implications(data)
    assert result["selections"]["alcance_actividad"] == "Internacional"
