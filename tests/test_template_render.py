"""Tests for template_handler.py — binary-copy + text-only rendering."""

import hashlib
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "templates" / "plantilla.docx"


@pytest.fixture(autouse=True)
def _require_template():
    if not TEMPLATE_PATH.exists():
        pytest.skip("plantilla.docx not found — skipping template tests")


def test_render_creates_output_file():
    """render_template() must return a path that exists."""
    from docx_generator.template_handler import render_template

    data = {
        "direct_fields": {"Razon_Social": "Test Corp", "CIF": "B12345678"},
        "selections": {},
        "free_texts": {},
    }
    output = render_template(document_id=9999, data=data, empresa_name="Test Corp")
    assert output.exists(), f"Output not created: {output}"
    output.unlink(missing_ok=True)


def test_template_is_not_modified():
    """The original plantilla.docx must be byte-identical before and after render_template()."""
    from docx_generator.template_handler import render_template

    before = hashlib.md5(TEMPLATE_PATH.read_bytes()).hexdigest()
    data = {"direct_fields": {}, "selections": {}, "free_texts": {}}
    output = render_template(document_id=9998, data=data)
    after = hashlib.md5(TEMPLATE_PATH.read_bytes()).hexdigest()
    assert before == after, "Template was modified by render_template()!"
    output.unlink(missing_ok=True)


def test_output_contains_company_name():
    """Direct field Razon_Social must appear somewhere in the rendered document text."""
    from docx import Document
    from docx_generator.template_handler import render_template

    data = {
        "direct_fields": {"Razon_Social": "EMPRESA_UNICA_XYZ"},
        "selections": {},
        "free_texts": {},
    }
    output = render_template(
        document_id=9997, data=data, empresa_name="EMPRESA UNICA XYZ"
    )
    doc = Document(str(output))
    all_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " " + cell.text
    assert "EMPRESA_UNICA_XYZ" in all_text
    output.unlink(missing_ok=True)


def test_no_green_background_applied():
    """After rendering, no cell should have the old green fill (70AD47)."""
    from docx import Document
    from docx_generator.template_handler import render_template
    from docx.oxml.ns import qn

    data = {
        "direct_fields": {},
        "selections": {"situacion_empresa": "Más de 2 años"},
        "free_texts": {},
    }
    output = render_template(document_id=9996, data=data)
    doc = Document(str(output))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    shd = tcPr.find(qn("w:shd"))
                    if shd is not None:
                        fill = shd.get(qn("w:fill"), "")
                        assert (
                            fill.upper() != "70AD47"
                        ), f"Old green fill found in cell: '{cell.text}'"
    output.unlink(missing_ok=True)
